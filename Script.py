###IMPORTATION DE MODULE
#module pour l'interface
import tkinter as tk
import tkinter.messagebox as msg
from tkinter import Label, Tk, Button, Frame, LabelFrame, filedialog, DoubleVar, Scale, HORIZONTAL

#module pour la gestion de fichier
import os
import urllib.request
from urllib.error import URLError
import ssl 
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#module pour les formules mathematiques
import math
from math import sqrt

#module pour la visualisation des données 
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from pandastable import Table
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure


###VARIABLES GLOBALES
PDB_fichier=[]
data=[]


###RECUPERATION DES FICHIERS

#via internet 
def afficher_input () :
    """Fonction qui génère une interface utilisateur permettant à l'utilisateur de saisir un code de fiche PDB"""
    #Génère la barre d'input 
    code_PDB = tk.Entry(fenetre_i)
    code_PDB.grid(row=8, column=0, columnspan=2)
    #Affiche les instructions supplémentaires
    script_input = "Veuillez saisir le code de la fiche PDB puis appuyer sur Valider."
    l = Label(fenetre_i, text= script_input, font=("Arial",8, "italic"), padx=5, pady=5)
    l.grid(row=7, column=0, columnspan=2, sticky='nsew')
    #Affiche le bouton valider qui exécute la recherche et lance la suite du code
    valider = tk.Button(fenetre_i, text=" Valider ", command= lambda : loadweb(code_PDB.get()))
    valider.grid(row=9, column=0, columnspan=2 )

def loadweb(codePDB):
    """ Fonction qui permet de charger le contenu d'une fiche PDB d'internet dans une liste.
    input: chaîne de caractères correspondant au codePDB
    output: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
            data, qui est une liste de listes. Chaque sous-liste est une ligne du fichier PDB. Chaque élément d'une sous liste est un mot de cette ligne"""
    try:#essaye de récupérer le fichier pdb sur internet à partir du code entré
        context = ssl._create_unverified_context() #autorisation pour certains sites
        file=urllib.request.urlopen(f"http://files.rcsb.org/view/{codePDB.upper()}.pdb", context=context)
        PDB_lines=file.readlines()  #essaye de le stocker dans une variable locale 
        file.close() 
    except URLError: #affiche une fenetre d'erreur pour indiquer que la recherche a échoué
        msg_error= ("Le code PDB est introuvable ou nous n'êtes pas connecté à internet. "+ "\n"*2
                    +"Vérifiez le code et votre connexion internet." + "\n"*2 + "Cliquez sur OK puis retapez votre code") # message d'erreur qui sera affiché
        msg.showerror("Code PDB Error \n", msg_error) #inialise la fenetre 
    else:
        #Récupère le code PDB dans une variable 
        global Code_PDB
        Code_PDB = codePDB.upper() #mise du code PDB en majuscules
        #stocke les valeurs dans des variables globales pour les utiliser dans la suite du code 
        global PDB_fichier
        global data
        PDB_fichier= []
        data=[]
        for ligne in PDB_lines:
            data.append(ligne.decode("utf8").strip()) #décode la ligne à partir d'une séquence d'octets en utilisant le jeu de caractères UTF-8 et enlève l'espace en début et en fin de chaine de caractère
            PDB_fichier.append(ligne.decode("utf8").strip()) 
        for i in range(len(data)):
            data[i]=data[i].split() #crée une liste de mots avec pour séparateur l'espace
        #Ferme la fenêtre d'initialisation
        fenetre_i.destroy()
        #Ouvre la fenetre d'exploitation
        app = Application()
        app.mainloop()

#en local
def ouvrir_fichier():
    """Fonction permettant à l'utilisateur de sélectionner un fichier PDB en utilisant une boîte de dialogue"""
    chemin_fichier = filedialog.askopenfilename(title="Sélectionner un fichier", filetypes=[("Fichiers PDB", "*.pdb")]) #ouvre la boite de dialogue permettant à l'utilisateur de sélectionner un fichier en .pdb
    nom_fichier = os.path.basename(chemin_fichier) # extrait le nom du fichier à partir du chemin complet du fichier sélectionné
    dir_path = os.path.dirname(chemin_fichier) # extrait le chemin du répertoire contenant le fichier sélectionné
    if dir_path != '' and nom_fichier != '': #si un fichier a été selectionné 
        os.chdir(dir_path) #normalise le chemin d'acces 
        loadlocal(dir_path, nom_fichier) #lance la suite des fonctions

def loadlocal (chemin_fichier, nom_fichier):
    """Fonction recherchant une fiche PDB stockée en local pour la stocker dans une liste.
    input: chaîne de caractère correspondant au  chemin d'accès local du fichier et chaine de caractère contenant le nom du fichier
    output: PDB_fichier et data, définis précédemment"""
    chemin_fichier= os.path.abspath(chemin_fichier) # standardise à partir d'un chemin absolu ou relatif le chemin d'accès à un fichier
    with open (nom_fichier, 'r') as file : # ouverture du fichier en mode lecture (l'utilisation de with open garantit que le fichier sera fermé automatiquement après utilisation)
        PDB_lines=file.readlines()
        #stocke les données dans des variables globales 
        global PDB_fichier
        global data
        PDB_fichier= []
        data=[]
    for ligne in PDB_lines:
        data.append(ligne.strip())
        PDB_fichier.append(ligne.strip())
    for i in range(len(data)):
        data[i]=data[i].split() # crée une liste de mot avec comme séparateur un espace
    #Récupère le code pbd sans l'extension
    global Code_PDB
    Code_PDB = (nom_fichier[:-4]).upper()
    #Ferme la fenêtre d'initialisation
    fenetre_i.destroy()
    #Ouvre la fenetre d'exploitation
    app = Application()
    app.mainloop()


###FENETRE D'INITIALISATION

def initialisation():
    """Fonction qui permet d'excuter d'autres fonction pour une créer fenêtre d'initialisation en utilisant la bibliothèque Tkinter"""
    #Ouvre la fenetre et défini la valeur en globale pour y faire appel dans d'autres fonctions
    global fenetre_i
    fenetre_i = Tk()
    # Attribue le nom "Recherche d'informations à partir d'un fichier PDB" à la fenêtre
    label = Label(fenetre_i, text=" Recherche d'informations à partir d'un fichier PDB",fg="blue", font=("Helvetica",12, "bold"), padx=5, pady=5)
    label.grid(row=1, column=0, columnspan=2, sticky='nsew')
    # Crée une sous fenetre "Instructions"
    l = LabelFrame(fenetre_i, text="Instructions", padx=5, pady=5)
    l.grid(row=4, column=0, columnspan=2, sticky='nsew')
    # Insère du texte dans la fenetre 
    script=("Ce programme permet d'extraire des informations d'un fichier PDB et de les visualiser."+ "\n" + 
             "Veuillez choisir le mode d'importation des données de la fiche PDB.")
    label_script = Label(l, text= script)
    label_script.grid(row=5, column=0, columnspan=2, sticky='nsew')

    #Chargement du fichier en ligne (création de bouton qui éxécute les fonctions précédentes)
    enligne = tk.Button(fenetre_i, text=" A partir d'un code PDB ", command=lambda: afficher_input()) # crée un bouton "A partir d'un code PDB" associé à la fonction "afficher_input()"
    enligne.grid(row=5, column=0, sticky='nsew')

    #Chargement du fichier en local (création de bouton qui éxécute les fonctions précédentes)
    local = tk.Button(fenetre_i, text=" A partir d'une fiche stockée en local", command=lambda: ouvrir_fichier()) # crée un bouton "A partir d'une fiche stockée en local" associé à la fonction "ouvrir_fichier()"
    local.grid(row=5, column=1, sticky='nsew')
    
    fenetre_i.mainloop() #Lancement de la fenetre tkinter

###RECUPERATION DE DONNEES 
def titre (PDB_fichier):
    """Fonction qui récupère le titre du fichier PDB.
    input: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
    output: chaine de caractère contenant le titre du fichier PDB"""
    titre= '' 
    PDB_court = PDB_fichier[: 15] #récupère seulement les premières lignes du fichier pour aller plus vite 
    if PDB_court[1].startswith ('TITLE'): #si la ligne 1 commence par titre 
        titre += (PDB_court[1][6:]) #recupère le nom du titre
        for i in range(len(PDB_court)) : #parcourt le reste du fichier à partir de la ligne 2 
            if (i+2) < len(PDB_court) and PDB_court[2+i].startswith ('TITLE'): #si la ligne suivant fait partie du titre aussi 
                titre += ("\n"+ PDB_court[i+2][7:]) #recupère le reste du titre sans le chiffre
            
    return titre

def nom_prot(PDB_fichier) : 
    """Fonction qui récupère le nom de la protéine étudiée dans le fichier PDB.
    input: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
    output: chaine de caractère contenant le nom de la protéine"""
    nom=""
    for ligne in PDB_fichier:
        if ligne.startswith ("COMPND") and "MOLECULE:" in ligne: #recherche le nom de la molécule 
            nom= (ligne[20:])
        if ";" in nom: #si le document pdb est issu d'un document txt il contient des ';' à la fin de chaque ligne
                nom= nom.replace(";", "") 
    return nom

def source(PDB_fichier):
    """Fonction qui récupère l'espèce à partir de laquelle la protéine est issue.
    input: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
    output: chaine de caractère contenant le nom de l'espèce à partir de laquelle la protéine est issue"""
    for ligne in PDB_fichier:
        if "SOURCE   2 ORGANISM_SCIENTIFIC" in ligne: # si "SOURCE   2 ORGANISM_SCIENTIFIC" est dans une des lignes du fichier PDB
            morceau_source = ligne.split(":") # sépare la ligne en utilisant ":" comme séparateur
            source = morceau_source[1]
            return source[:-1]
    return None # si aucune ligne ne correspond à ces conditions, la fonction renvoit None

def methode(data):
    """Fonction qui récupère la méthode de résolution de la structure protéique.
    input: data, qui est une liste de listes. Chaque sous-liste est une ligne du fichier PDB. Chaque élément d'une sous liste est un mot de cette ligne
    output: chaine de caractère contenant le nom de la méthode de résolution de la structure protéique"""
    liste_methode = []
    liste_resolution=[]
    for i in data:
        if i[0] == "EXPDTA": # si le premier élément d'une ligne est EXPDTA
            liste_methode.append(i[1:]) # ajoute cette ligne à la liste liste_method, sans "EXPDTA"
    methode_resolution = ' '.join([" ".join(line) for line in liste_methode]) # transforme la liste en chaine de caractère
    return methode_resolution

def resolution (PDB_fichier):
    """Fonction qui récupère la résolution de la structure protéique.
    input: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
    output: chaine de caractère contenant la valeur de la résolution de la méthode d'analyse de la structure protéique"""
    for ligne in PDB_fichier:
        if "REMARK   2 RESOLUTION" in ligne: # si une des lignes du fichier contient "REMARK   2 RESOLUTION"
            morceau_resolution= ligne.split(".") # sépare la ligne en utilisant "." comme séparateur
            resolution= morceau_resolution[1]+"."+ morceau_resolution[2]
            return resolution[:-1]
    return "il n'y a pas de résolution pour cette méthode" # si aucune ligne ne correspond à ces conditions c'est qu'il n'y a pas de résolution pour cette méthode

def header(PDB_fichier):
    """Fonction qui récupère les informations du header de la structure protéique.
    input: PDB_fichier, qui est une liste contenant toutes les lignes d'un fichier PDB
    output: chaine de caractère contenant les informations du header"""
    PDB_head= PDB_fichier[0: 250] #récupère les premières lignes du fichier pour aller plus vite
    version = ''
    chain = ''
    esp= nom_prot(PDB_head)
    species= ''
    nume = ''
    for ligne in PDB_head: #parcourt le document court 
        if ligne.startswith('SOURCE') and 'MOL_ID:' in ligne: #recherche le nombre de molécule décrite
            version= ligne[17:]
            if ";" in version: #si le document pdb est issu d'un document txt il contient des ';' à la fin de chaque ligne
                version= version.replace(";", "") 
    
        if ligne.startswith ("COMPND") and "CHAIN" in ligne:#recherche l'identifiant de la chaine
            chain= ("Chain" + str(ligne[17:]))
            if ";" in chain : 
                chain= chain.replace(";", "")
    
        if ligne.startswith ("SOURCE") and "ORGANISM_SCIENTIFIC:" in ligne: #recherhce le nom de l'organisme scientifique source 
            species= str(ligne[31:])
            if ";" in species : 
                species= species.replace(";", "")
    
        if ligne.startswith ("SOURCE") and "ORGANISM_TAXID:" in ligne: #identifiant de l'organisme scientifique source 
            nume= str(ligne[27:])
            if ";" in nume : 
                nume= nume.replace(";", "")

    #attribue ces valeurs à head avec la nomenclature d'un header 
    head = (">" + Code_PDB + str(version) + "|" + str(chain)+ "|" + str(esp) + "|" + str((species[1:].lower()).capitalize()) +" (" + str(nume) + ")" + "\n")
    return head 

def Calpha (data):
    """Fonction qui récupère les lignes d'une fiche PDB contenant des carbones alpha,
    input: data, qui est une liste de listes. Chaque sous-liste est une ligne du fichier PDB. Chaque élément d'une sous liste est un mot de cette ligne
    output: liste des lignes de chaque carbone alpha"""
    Calpha=[]
    for ligne in data:
        if ligne[0]=='ATOM' and 'CA' in ligne: # si une ligne commence par "ATOM" et contient "CA"
                Calpha.append(ligne) # récupération de la ligne dans la liste Calpha
    return Calpha

def taille (Calpha):
    """Fonction renvoyant le nombre d'acides aminés résolus dans la séquence
    input : liste contenant les lignes d'une fiche PDB qui contennient des carbones alpha
    output : nombre d'acides aminés résolus"""
    taille_prot=len(Calpha) # le nombre d'acides aminés dans la séquence correspond au nombre d'éléments de la liste Calpha
    return taille_prot

#Défintion de variables globales à l'aide d'une fonction 
def definition_var (): 
    global atome
    atome=Calpha (data)
    global sequence_3Le
    sequence_3Le=sequence_3L(atome)
    global sequence_1Le
    sequence_1Le=sequence_1L(sequence_3Le)
    global seq
    seq=sequence_aa(sequence_1Le)


###CREATION DE LA SEQUENCE FASTA

def sequence_3L (Calpha):
    """Fonction renvoyant la séquence protéique d'un fichier PDB, séquence codée avec le code à 3 lettres
    input: liste contenant les lignes d'une fiche PDB qui contennient des carbones alpha
    output: liste dont chaque élément est un acide aminé de la séquence, écrit avec le code à 3 lettres"""
    sequence_3L=[]
    for aa in Calpha:
        sequence_3L.append(aa[3][-3:]) #prend les 3 dernières valeurs au cas ou l'acide aminés contiendrait 4 lettres
    return sequence_3L

def sequence_1L (sequence_3Lettre):
    """ Fonction qui traduit une séquence protéique à 3 lettres en séquence protéique à 1 lettre
    #input: liste des acides aminés codés avec 3 lettres
    #output: liste des acides aminés codés avec 1 lettre """
    dico_3L_1L={
    'ALA': 'A', 'ARG': 'R', 'ASN': 'N', 'ASP': 'D', 'CYS': 'C', 'GLN': 'Q', 'GLU': 'E',
    'GLY': 'G', 'HIS': 'H', 'ILE': 'I', 'LEU': 'L','LYS': 'K', 'MET': 'M', 'PHE': 'F',
    'PRO': 'P', 'SER': 'S', 'THR': 'T', 'TRP': 'W','TYR': 'Y','VAL': 'V'} # dictionnaire faisant le lien entre le code à 3 lettres et le code à une lettre
    sequence_1Lettres=[]
    for i in sequence_3Lettre:
        sequence_1Lettres.append(dico_3L_1L[i])
    return sequence_1Lettres

def sequence_aa(sequence_1Lettre):
    """Fonction qui ajoute un retour à la ligne tous les 80 acides aminés
    input: liste des acides aminés à une lettre
    output: chaine de caractère avec les acides aminés à une lettre et un retour à la ligne tous les 80 caractères"""
    ligne_aa = ''.join(sequence_1Lettre)
    lines = [ligne_aa[i:i + 80] for i in range(0, len(ligne_aa), 80)] # crée des segments de 80 acides aminés
    return '\n'.join(lines) # entre chaque segment de 80 acides aminés, ajout d'un retour à la ligne    


###PAGE DE DESCRIPTION

def affichage_informations ():
    #Récupération des données
    titre_prot=titre(PDB_fichier)
    global atome
    atome=Calpha (data)
    taille_prot=taille(atome)
    espece= source(PDB_fichier)
    nom = nom_prot(PDB_fichier)
    methode_resolution=methode(data)
    valeur_resolution=resolution(PDB_fichier)
    #sequence fasta:
    global sequence_3Le
    sequence_3Le=sequence_3L(atome)
    global sequence_1Le
    sequence_1Le=sequence_1L(sequence_3Le)
    global seq
    seq=sequence_aa(sequence_1Le)
        
    script = ((f"Le titre est{titre_prot}")+ "\n"*2 + f"Le nom de la protéine est {nom}." + "\n" + (f"Cette protéine provient de l'espèce {espece}.") +"\n"*2 +
              (f"Elle a une taille {taille_prot} acides aminés.")+ "\n" + ("Sa séquence en acides aminés est :")+
              "\n" + seq + "\n"*2 +(f'La méthode de résolution utilisée est {methode_resolution}.')+ "\n" +
              (f'La résolution est de{valeur_resolution}') )
    return script              

#Enregistrer au format Fasta
def fenetre_enr(): 
    """ Fonction qui permet d'enregistrer la fiche fasta créée avec les données du fichier PDB"""
    #ouverture du gestionnaire de fichier pour déterminer le chemin d'enregistrement du fichier txt
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", initialfile=Code_PDB)
    if file_path != '' : #si un chemin est selectionné
        with open(file_path, 'w') as f:
            f.write(enregistrer_fichier(seq)) #cré le fichier txt

def enregistrer_fichier(sequence):
    """ Fonction qui permet de rédiger le contenu d'une fiche fasta
    input : la sequence d'acides aminés à 1 lettre ordonnée à 80 caractères par ligne
    output: une chaine de caractère comprenant le header et la séquence à 1 Lettre"""
    Header = header(PDB_fichier)
    file = Header
    file += sequence
    return str(file)

#Enregistrer toutes les données générées 
def enregistrer_final(): 
    """ Fonction qui permet de générer un fichier .docx contenant toutes les données générées lors de cette analyse ; y compris les tableaux et graphiques """
    #ouverture du gestionnaire de fichier pour le choix du dossier d'enregistrement
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"Analyse_{Code_PDB}")
    if file_path != '' : #si le chemin est selectionné
        doc = Document() # Créer un nouveau document Word        
        
        # Définir les styles de paragraphe
        style = doc.styles.add_style("Titre_1",1)
        style.font.name = 'Arial' #Police
        style.font.size = Pt(14) #taille
        style.font.bold = True #en gras 
        style.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) #couleur
        
        style1 = doc.styles.add_style('Titre_p',1)
        style1.font.name = 'Arial'
        style1.font.size = Pt(12)
        style1.font.bold = True
        style1.font.color.rgb = RGBColor(0x07, 0x60, 0xF9)
        
        style2 = doc.styles.add_style('Basic',1)
        style2.font.name = 'Arial'
        style2.font.size = Pt(11)
        paragraph_format2 = style2.paragraph_format
        paragraph_format2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #type d'alignement

        ##Titre document 
        titre = doc.add_heading(f'Analyse de la fiche PDB {Code_PDB.upper()}', level=0)
        titre.style = style
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        doc.add_paragraph() #saut de ligne

        ##Description du document pdb 
        titre1 = doc.add_heading("Description de la fiche", level=1)
        titre1.style = style1
        cell_1 = doc.add_paragraph(affichage_informations ())
        cell_1.style = style2
        doc.add_paragraph()

        ##Composition en AA du document pdb 
        titre2 = doc.add_heading("Composition en acides aminés", level=1)
        titre2.style = style1
        script_compo = ("Une analyse de la composition en acides aminés de la séquence étudiée a été réalisé. Les résultats sont répertoriés ci-dessous." + 
                        "\n"*2 +  f"La séquence contient {len(seq)} acides aminés." + "\n" + seq)
        compo = doc.add_paragraph(script_compo)
        compo.style = style2
        compo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        #Importer les valeurs pour la suite des analyses 
        definition_var ()
        affichage_stat() 
        frequence_moyenne_acides_aminés = {'A': 8.25, 'R': 5.53, 'N': 4.06, 'D': 5.46, 'C': 1.38, 'Q': 3.93, 'E': 6.72, 'G': 7.07, 'H': 2.27, 'I': 5.91,
                                               'L': 9.65, 'K': 5.80, 'M': 2.41, 'F': 3.86, 'P': 4.74, 'S': 6.65, 'T': 5.36, 'W': 1.10, 'Y': 2.92, 'V': 6.85}
        #création d'un tableau comparatif des fréquences 
        df = pd.DataFrame({"AcideAminé": list(Nombre_AA.keys()), 'Fréquence Observée': list(Nombre_AA.values()), 'Fréquence Moyenne': list(frequence_moyenne_acides_aminés.values())})

        #Création d'un tableau dans le document .docx en suivant la trame de df
        table2 = doc.add_table(df.shape[0]+1, df.shape[1])
        for j in range(df.shape[-1]): # Ajouter les en-têtes de la table
            table2.cell(0,j).text = df.columns[j]
        for i in range(df.shape[0]): # Ajouter les données de la table
            for j in range(df.shape[-1]):
                table2.cell(i+1,j).text = str(df.values[i,j])
        
        #Importation du graphique en l'enregistrant en .png 
        fig=graphique_composition(Nombre_AA) # création du graphique
        fig.savefig('figure.png') #enregistrement en .png 
        doc.add_picture('figure.png', width=Inches(4.0))
        doc.add_paragraph()

        ##Profil d'hydrophobicite
        titre3 = doc.add_heading("Profil d'hydrophobicité", level=1)
        titre3.style = style1
        script_hydro = ("Une analyse du profil d'hydrophobicité de la séquence étudiée a été réalisé. Les résultats sont répertoriés ci-dessous." + 
                        "Le profil d'hydrophobicité des acides aminés de la séquence a été calculé grâce à l'échelle d'hydrophobicité de Fauchere et Pliska sur une fenêtre glissante de 9 acides aminés.")
        hydro = doc.add_paragraph(script_hydro)
        hydro.style = style2
        hydro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        window_size = DoubleVar(value=9) #initialisation de la fenetre de lecture
        position_med,hydrophobicite_moy = profil_hydrophobicite(sequence_1Le,window_size)
        tab_hydro = tableau_profile_hydrophobicite(position_med, hydrophobicite_moy) #création du tableau trame pour l'hydrophobicité
        
        #Création d'un tableau dans le document .docx en suivant la trame de tab_hydro
        table3 = doc.add_table(tab_hydro.shape[0]+1, tab_hydro.shape[1])
        for j in range(tab_hydro.shape[-1]): # Ajouter les en-têtes de la table
            table3.cell(0,j).text = tab_hydro.columns[j]
        for i in range(tab_hydro.shape[0]): # Ajouter les données de la table
            for j in range(tab_hydro.shape[-1]):
                table3.cell(i+1,j).text = str(tab_hydro.values[i,j])
        
        grap_hydro=graphique_profile_hydrophobicite(position_med, hydrophobicite_moy) #création du graphique pour l'hydrophobicite 
        grap_hydro.savefig('figure_hydro.png') #sauvegarde en .png
        doc.add_picture('figure_hydro.png', width=Inches(4.0))
        doc.add_paragraph()

        ##Pont disulfure
        titre4 = doc.add_heading("Prédiction de ponts disulfures", level=1)
        titre4.style = style1
        resul_sulf = affichage_stru() #analyse de la présence de ponts 
        script_sulf = ("Une analyse de la structure de la séquence a permis la prédiction de ponts disulfures intramoléculaires. Les résultats sont répertoriés ci-dessous." + 
                    "\n"+ "NB : si les cystines présentes ne sont pas impliqués dans la structure elle pourrait participer à un pont disulfure avec une autre molécule (non étudiée ici)." + 
                    "\n"*2 + resul_sulf) #texte de l'analyse
        sulf = doc.add_paragraph(script_sulf)
        sulf.style = style2
        sulf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc.add_paragraph()

        ##Matrice contact
        titre5 = doc.add_heading("Matrice de contact", level=1)
        titre5.style = style1
        doc.add_paragraph()
        script_ma =(f"La matrice de contact de la protéine correspondante au code PDB {Code_PDB} est présentée ci-dessous" + 
                         "Une matrice de contact de protéine est une représentation matricielle des interactions entre les résidus"+ 
                         " d’une protéine tridimensionnelle." +" Elle indique quels résidus sont en contact et quels résidus sont éloignés selon une échelle de coloration." + 
                         " Elle permet de visualiser la structure secondaire et tertiaire de la protéine, ainsi que les domaines et les motifs structuraux.")
        matrice = doc.add_paragraph(script_ma)
        matrice.style = style2
        matrice.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        CALPHA = Calpha (data)
        matrice_contact = calcul_distance_calpha(CALPHA)
        m=afficher_matrice_contact(matrice_contact) #création de la matrice 
        m.savefig('matrice_contact.png') #enregistrement en png
        doc.add_picture('matrice_contact.png', width=Inches(4.0))

        # Enregistrer le document
        doc.save(file_path)


###PAGE DE COMPOSITION
def enregistrer_tableau (df, nom): 
    """ Fonction qui permet d'enregistrer un tableau sous le format excel 
    input : tableau pandas
    output : fichier excel enregistrer dans le chemin demandé avec les valeurs de la table d'input """
    #ouverture du gestionnaire de fichier pour déterminer le chemin d'enregistrement du fichier .xlsx
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", initialfile=(f"{nom}{Code_PDB}"))
    if file_path != '' : #si un chemin est selectionné
        df.to_excel(file_path) #transforme le tableau source en excel 

def analyse_composition(sequence_1Lettre): 
    """Fonction qui permet l'analyse de fréquence des acides aminés dans une séquence,
    input: liste de caractères correspondant à la séquence en acides aminés codés avec 1 lettre
    output: dictionnaire où la clé est le code de l'acide aminé à une lettre, et la valeur associée est la fréquence"""
    frequence_moyenne_acides_aminés = {'A': 8.25, 'R': 5.53, 'N': 4.06, 'D': 5.46, 'C': 1.38, 'Q': 3.93, 'E': 6.72, 'G': 7.07, 'H': 2.27, 'I': 5.91,
                                       'L': 9.65, 'K': 5.80, 'M': 2.41, 'F': 3.86, 'P': 4.74, 'S': 6.65, 'T': 5.36, 'W': 1.10, 'Y': 2.92, 'V': 6.85}
    compteur = {aa: 0 for aa in frequence_moyenne_acides_aminés}  # initialise un dictionnaire pour chaque acide aminé avec une fréquence de 0
    for aa in sequence_1Lettre:
        compteur[aa] += 1  # compte le nombre d'occurrences de chaque acide aminé dans la séquence
    total_aa = sum(compteur.values())  # calcule le nombre total d'acides aminés dans la séquence
    for aa in compteur:
        compteur[aa] = (compteur[aa] / total_aa) * 100  # calcule la fréquence en pourcentage de chaque acide aminé
    for aa in compteur:  # comparaison avec les fréquences moyennes de la fiche swissprot
        if aa in frequence_moyenne_acides_aminés:
            diff = compteur[aa] - frequence_moyenne_acides_aminés[aa]
    return compteur, diff

def analyse_composition_3L(sequence_3Lettre): 
    """Fonction qui permet l'analyse de fréquence des acides aminés dans une séquence,
    input: liste de caractères correspondant à la séquence en acides aminés codés avec 3 lettres
    output: dictionnaire où la clé est le code de l'acide aminé à une lettre, et la valeur associée est la fréquence"""
    frequence_moyenne_acides_aminés = {'ALA': 8.25, 'ARG': 5.53, 'ASN': 4.06, 'ASP': 5.46, 'CYS': 1.38, 'GLN': 3.93, 'GLU': 6.72, 'GLY': 7.07, 'HIS': 2.27, 'ILE': 5.91,
                                       'LEU': 9.65, 'LYS': 5.80, 'MET': 2.41, 'PHE': 3.86, 'PRO': 4.74, 'SER': 6.65, 'THR': 5.36, 'TRP': 1.10, 'TYR': 2.92, 'VAL': 6.85}
    compteur = {aa: 0 for aa in frequence_moyenne_acides_aminés}  # initialise un dictionnaire pour chaque acide aminé avec une fréquence de 0
    for aa in sequence_3Lettre:
        compteur[aa] += 1  # compte le nombre d'occurrences de chaque acide aminé dans la séquence
    total_aa = sum(compteur.values())  # calcule le nombre total d'acides aminés dans la séquence
    for aa in compteur:
        compteur[aa] = (compteur[aa] / total_aa) * 100  # calcule la fréquence en pourcentage de chaque acide aminé
    return compteur

def graphique_composition (compteur):
    """ Fonction qui crée un barplot de la composition en acides aminés de la séquence étudiée 
    input : dictionnaire où la clé est le code de l'acide aminé à une lettre, et la valeur associée est la fréquence
    output : graphique matplotlib de la composition en acides aminés """
    frequence_moyenne_acides_aminés = {'A': 8.25, 'R': 5.53, 'N': 4.06, 'D': 5.46, 'C': 1.38, 'Q': 3.93, 'E': 6.72, 'G': 7.07, 'H': 2.27, 'I': 5.91,
                                       'L': 9.65, 'K': 5.80, 'M': 2.41, 'F': 3.86, 'P': 4.74, 'S': 6.65, 'T': 5.36, 'W': 1.10, 'Y': 2.92, 'V': 6.85}
    titre = list(compteur.keys()) # Liste des acides aminés
    frequencies = list(compteur.values()) # Liste des fréquences observées
    moyennes = list(frequence_moyenne_acides_aminés.values()) # Liste des fréquences moyennes
    x = np.arange(len(titre)) # Création d'un tableau d'indices pour les étiquettes
    width = 0.3 # Largeur des barres dans le graphique
    fig, ax = plt.subplots(figsize=(5, 4))
    rects1 = ax.bar(x - width/2, frequencies, width, label='Fréquence observée') # Création des barres pour les fréquences observées
    rects2 = ax.bar(x + width/2, moyennes, width, label='Fréquence moyenne') # Création des barres pour les fréquences moyennes
    ax.set_ylabel('Fréquence en pourcentage')
    ax.set_title(("Comparaison des fréquences observées" + "\n" +" par rapport aux fréquences moyennes"), fontsize= 12, fontweight = "bold")
    ax.set_xticks(x)
    ax.set_xticklabels(titre)
    ax.legend()
    return fig 

def tableau_composition (compteur):
    """ Fonction qui permet de créer un tableau des fréquences en acides aminés dans la séquence étudiée à partir d'un dictionnaire
    input : dictionnaire où la clé est le code de l'acide aminé à une lettre, et la valeur associée est la fréquence
    output : table de la composition en acides aminés sous le format pandas"""
    frequence_moyenne_acides_aminés = {'A': 8.25, 'R': 5.53, 'N': 4.06, 'D': 5.46, 'C': 1.38, 'Q': 3.93, 'E': 6.72, 'G': 7.07, 'H': 2.27, 'I': 5.91,
                                       'L': 9.65, 'K': 5.80, 'M': 2.41, 'F': 3.86, 'P': 4.74, 'S': 6.65, 'T': 5.36, 'W': 1.10, 'Y': 2.92, 'V': 6.85}
    # Création du DataFrame pour le tableau comparatif
    df = pd.DataFrame({"AcideAminé": list(compteur.keys()), 'Fréquence Observée': list(compteur.values()), 'Fréquence Moyenne': list(frequence_moyenne_acides_aminés.values())})
    return df 

def affichage_stat():
    """Fonction qui permet la création des variables nécéssaires pour l'analyse de composition en acides aminés"""
    definition_var()
    #fréquence aa :
    global Nombre_AA
    Nombre_AA, _= analyse_composition(sequence_1Le)


##PAGE DE STRUCTURE
def affichage_stru():
    """ Fonction qui permet la  création du renvoie des analyses de présence de ponts disulfures dans la séquence étudiée en utilisant des varibales globales 
    output : chaine de caractères comprenant les résultats de l'analyses  """
    #initialisation des variables nécessaires à l'analyse grâce aux différentes fonctions reliées
    dico_soufre=atome_soufre(data)
    liste_distance, liste_positions, liste_souffre_implique=pont_disulfure (dico_soufre)
    liste_non_implique=soufres_non_impliques(dico_soufre, liste_souffre_implique)
    
    pont_hypothetique = ''
    aa_non_implique = ''
    
    if len(dico_soufre) == 0 : # s'il n'y a pas de cystéine
        script = ("La structure étudiée ne contient pas d'atome de cystéine."+ "\n"+" Il ne peut donc pas y avoir de ponts disulfures dans la structure.")
        return script
    
    if liste_distance != '': # s'il y a des ponts disulfures possibles
        text1 = "Si la protéine est sécrété, un pont disulfure pourrait voir le jour entre les cystéines en positions "
        for positions, distance in zip(liste_positions, liste_distance):
            pont_hypothetique += (f' {positions} car leur distance est de {distance} Å' + "\n")
        text4=(text1 + "\n" + pont_hypothetique)
    else : 
        text4 = ''
    
    if len(liste_non_implique) != 0: # s'il y a des cys non impliquées dans les ponts 
        if len(liste_non_implique) == len (dico_soufre): #si toutes les cystéines ne sont pas impliquées 
            script = (f"La séquence étudiée contient {len(dico_soufre)} cystéine(s). " + "\n" + 
                     "Cependant aucune cystéine n'est impliquée dans la formation de ponts disulfures,"+"\n"+
                     " car les distances qui les sépares sont supérieures à 3 Å.")
            return script
        else : # si quelques cystéines ne sont pas impliquées 
            text2 = (f'Certains soufres ne pourraient pas être impliqués dans des ponts disulfures, voici leurs positions:')
            for cys in liste_non_implique : 
                aa_non_implique += (cys+ "\n")
            text3 = (text2 + "\n" + aa_non_implique)
    else : # s'il n'y a pas de cys non impliquées
        text3 = ''
    script = (text4 + "\n"*2 +  text3)
    return script 

def atome_soufre(data):
    """Fonction qui permet d'obtenir les positions des atomes de soufres et leurs coordonnés en plan x,y,z
    input: une liste contenant toutes les lignes du codes contenant des sous listes pour chaque mots
    output: dictionnaire avec comme clé le nombre de l'atome dans la séquence et comme valeur les coordonnées en x,y,z"""
    dico_cysteine = {}
    for i in data:
        if i[0] == "ATOM" and i[2] == "SG": 
            numero_residu = i[5]
            dico_cysteine[numero_residu] = [i[6], i[7], i[8]]
    return dico_cysteine

def pont_disulfure(dico_cysteine):
    """ Fonction qui permet d'obtenir les possibles ponts disulfures entre deux cystéines ainsi que leurs distances tout en prennat en compte qu'un cystéine ne peut réaliser q'un pont disulfure
    input: dictionnaire avec comme clé le nombre de l'atome dans la séquence et comme valeur les coordonnées en x, y, z
    output: 3 listes : liste_distance = contenant les distances entres les atomes qui peuvent être impliqués dans un pont disulfures
    liste_positions = liste des positions des numéros des ponts disulfures qui sont possiblement impliqués dans un pont disulfure
    liste_cysteine_implique = liste des positions des cystéines qui peuvent êtres impliqué dans un pont disulfure"""
    liste_positions = []
    liste_distance = []
    liste_cysteine_implique = [] # liste incrémentée à chaque fois qu'un soufre apparait et donc une cysteine participe à un pont disulfure
    for i in dico_cysteine:
        for j in dico_cysteine:
            if i != j and i not in liste_cysteine_implique and j not in liste_cysteine_implique: # vérifie que les soufres sélectionnés ne sont pas déjà inclus dans un pont disulfures
                distance = sqrt((float(dico_cysteine[i][0]) - float(dico_cysteine[j][0]))**2 +
                                (float(dico_cysteine[i][1]) - float(dico_cysteine[j][1]))**2 +
                                (float(dico_cysteine[i][2]) - float(dico_cysteine[j][2]))**2) # calcul de la distance euclidienne entre deux soufres
                if distance <= 3: # on admet la possibilité d'un pont disulfure seulement si la distance est inférieur à 3
                    liste_distance.append(round(distance, 3)) # arrondi au milième la distance euclidienne entre les deux soufres
                    liste_positions.append(i + "-" + j)
                    liste_cysteine_implique.append(i) # liste incrémentée à chaque fois qu'un soufre participe à un pont disulfure
                    liste_cysteine_implique.append(j)
    return liste_distance, liste_positions, liste_cysteine_implique

def soufres_non_impliques(dico_cysteine, liste_cysteine_implique):
    """Fonction qui permet d'obtenir la position des cysteines qui ne sont pas impliqués dans des ponts disulfures
    input: dictionnaire avec comme clé le nombre de l'atome dans la séquence et comme valeur les coordonnées en x,y,z et la liste des cysteines qui peuvent êtres impliqué dans un pont disulfure
    output: une liste qui contient la position des cysteines qui ne sont pas impliqués dans un pont disulfure"""
    ensemble_cysteine= list(dico_cysteine.keys())# on obtient une liste de toutes les positions des soufres
    liste_cysteine_non_impliques = [cysteine for cysteine in ensemble_cysteine if cysteine not in liste_cysteine_implique] #créee une liste avec les soufres ne participant pas à des ponts disulfures
    return liste_cysteine_non_impliques


###PAGE de PHYSICO-CHIMIE
def classification(factor,Classification):
    """Fonction qui remplace la valeur du b-factor selon la classification choisie
    input : liste comprenant les informations des lignes atomes à partir du nom de l'atome
    output : valeur du b-factor et valeur de la nouvelle classification"""
    if Classification == 2:  #polarite
        #Création du dictionnaire spécifique
        polaires_non_charges = ('SER', 'THR', 'ASN', 'GLN', 'CYS')
        polaires_acides =  ('ASP', 'GLU')
        polaires_basiques = ('LYS', 'ARG', 'HIS')
        apolaires_non_aromatiques = ('GLY', 'ALA', 'VAL', 'LEU', 'ILE', 'PRO','MET')
        apolaires_aromatiques = ('PHE', 'TYR', 'TRP')
        groupe_polarite = {polaires_non_charges: 1, polaires_acides: 200, polaires_basiques: 400, apolaires_non_aromatiques: 600, apolaires_aromatiques : 800}
        #Attribution de la nouvelle valeur selon les lignes factors sélectionnées
        for element in groupe_polarite:
            if factor[0] in element:
                old_p = factor[7]
                new= round(groupe_polarite[element],2)
        return old_p, new
    
    elif Classification == 3: #poids moleculaire
        #Création du dictionnaire spécifique
        poids_moleculaires_aa = {'ALA': 89.09,'ARG': 174.20,'ASN': 132.12,'ASP': 133.10,'CYS': 121.15,
                                 'GLN': 146.15, 'GLU': 147.13, 'GLY': 75.07, 'HIS': 155.16, 'ILE': 131.18,
                                 'LEU': 131.18, 'LYS': 146.19, 'MET': 149.21, 'PHE': 165.19, 'PRO': 115.13, 
                                 'SER': 105.09,'THR': 119.12, 'TRP': 204.23, 'TYR': 181.19, 'VAL': 117.15}
        for element in poids_moleculaires_aa : 
            poids_moleculaires_aa[element]*4.5 #pour étendre les valeurs jusqu'à 999
        #Attribution de la nouvelle valeur selon les lignes factors sélectionnées 
        for element in poids_moleculaires_aa:
            if factor [0] in element:
                old_m = factor[7]
                new= round(poids_moleculaires_aa[element],2)
        return old_m, new

    elif Classification == 4: #frequence
        #Création du dictionnaire spécifique
        definition_var()
        dico_freq= analyse_composition_3L(sequence_3Le)
        for cle in dico_freq:
            dico_freq[cle] = dico_freq[cle] *50 #pour étendre les valeurs de fréquence
        #Attribution de la nouvelle valeur selon les lignes factors sélectionnées
        for element in dico_freq:
            if factor[0] in element:
                old_f = factor[7]
                new= round(dico_freq[element],2)
        return old_f, new
    
    elif Classification == 1 : #b-factor
        return factor[7], factor[7]

def fenetre_enr_hydro(Classification): 
    type_coloration = {1: "B_factor", 2: "Polarite", 3: "Poids_moleculaire", 4:"Frequence"} #traduit le code pour le nom du fichier 
    #ouverture du gestionnaire de fichier pour déterminer le chemin d'enregistrement du fichier txt
    file_path = filedialog.asksaveasfilename(defaultextension=".pdb", initialfile=(f"{Code_PDB}_{type_coloration.get(Classification)}"))
    if file_path != '' : # si un chemin a été selectionné 
        new_PDB= PDB_fichier
        with open(file_path, 'w') as f:
            for ligne in new_PDB:
                if ligne.startswith('ATOM'): #si l'element est une ligne d'informations sur des atomes 
                    ligne_separe= ligne.split() #sépare l'élément en liste de caractères
                    factor = ligne_separe[3:11] # récupère le nom jusqu'au b-factor
                    old, new = classification(factor,Classification) #retourne le b-factor et la valeur induite par la classification
                    indice = ligne.index(old)  # trouver l'indice de l'élément à remplacer
                    new_ligne = (ligne[:indice] + str(new) + ligne[indice + len(old) :]) # remplacer l'élément
                    
                    f.write(new_ligne)
                    f.write("\n")
                else: #si l'élément ne correspond pas à une ligne d'informations sur les atomes 
                    f.write(ligne)
                    f.write("\n")
        global path_hydro
        path_hydro = file_path


###PAGE MATRICE DE CONTACT
def calcul_distance_calpha(Calpha):
    """Fonction qui calcule la distance entre tous les carbones alpha de la protéine.
    input: liste des lignes contenant les carbones alpha
    output: matrice des distances entre les carbones alpha"""
    coord_calpha = []  # initialisation de la liste pour stocker les coordonnées des carbones alpha
    for ligne in Calpha: # extraction des coordonnées des carbones alpha
        if ligne[0] == 'ATOM' and ligne[2] == 'CA':  # vérifie si la ligne correspond à un atome de carbone alpha
            x = float(ligne[6])
            y = float(ligne[7])
            z = float(ligne[8])
            coord_calpha.append([x, y, z])  # ajoute les coordonnées à la liste
    n = len(coord_calpha)
    distances = np.zeros((n, n))
    # Calcul de la distance entre les carbones alpha
    for i in range(n): # parcours des indices des carbones alpha de 0 à n-1 -> sélection du carbone alpha de référence
        for j in range(i + 1, n): # parcours des indices des carbones alpha à partir de i+1 jusqu'à n-1 -> sélection du carbone alpha distinct de celui de référence
            distance = np.linalg.norm(np.array(coord_calpha[i]) - np.array(coord_calpha[j])) # calcul de la distance entre chaque paire de carbones alpha
            distances[i, j] = distance
            distances[j, i] = distance  # la matrice est symétrique
    return distances

def afficher_matrice_contact(matrice_contact):
    """Affiche la matrice de contact sur un graphique
    input: matrice de contact
    output : graph de la matrice de contact"""
    # Création de la figure et des axes
    fig = Figure(figsize=(5, 4), dpi=100)
    ax = fig.add_subplot(111)
    # Création du graphique de la matrice de contact
    cax = ax.imshow(matrice_contact, cmap='jet', interpolation='nearest')
    # Ajoute une barre de couleur à côté du graphique pour représenter la distance
    fig.colorbar(cax, label='Distance')
    ax.set_title('Matrice de Contact')
    ax.set_xlabel('Amino acid index')
    ax.set_ylabel('Amino acid index')
    return fig

###PAGE DU PROFIL D'HYDROPHOBICITE

def profil_hydrophobicite (sequence_1L, window_size):
    """Fonction qui calcule l'hydrophobicité moyenne pour chaque fenêtre glissante d'une longueur en acides aminées données par l'utilisateur
    input: liste contenant la séquence d'acides aminés sous la forme du code à 1 lettre
    output:deux listes :
                        position_mediane : liste avec les positions (entier) des acides aminés dans la séquence qui sont reliés à des profils d'hydrophobicité
                        hydrophobicite_moyenne: liste avec les floattants des valeurs  d'hydrophobicité moyennes pour chaque fenetre glissante"""
    hydrophobicity_scale = {'A': 0.31, 'R': -1.01, 'N': -0.60, 'D': -0.77, 'C': 1.54,'E': -0.64, 'Q': -0.22, 'G': 0.00, 'H': -0.96, 'I': 1.8,'L': 1.7, 'K': -0.99, 'M': 1.23, 'F': 1.79, 'P': 0.72,'S': -0.04, 'T': 0.26, 'W': 2.25, 'Y': 0.96, 'V': 1.22} #selon Fauchere & Pliska
    hydrophobicite_moyenne = []
    position_mediane=[]
    valeur_window_size= int(window_size.get()) #convertie la valeur de l'input en int pour pouvoir l'utiliser
    for i in range(len(sequence_1L) - valeur_window_size + 1):# +1 inclus la dernier acides aminés dans l'itération
        window = sequence_1L[i:i+valeur_window_size] # sélection des acides aminés dans la fenêtre glissante
        window_hydrophobicity = [hydrophobicity_scale[aa] for aa in window] # création d'une liste contenant les valeurs d'hydrophobicitées correspondant aux acides aminés présents dans la fenêtre glissante
        hydrophobicite_moyenne.append(round(sum(window_hydrophobicity) / valeur_window_size,3)) # réalise la moyenne des valeurs d'hydrophobicitées des acides aminés de la fenêtre glissante arrondi au millième
        position_mediane.append(math.ceil(((2*i)+ valeur_window_size/2))) # prend la moyenne des bornes de la fenêtre glissante et arrondi à l'entier supérieure si la valeur est un flottant
    return position_mediane, hydrophobicite_moyenne

def tableau_profile_hydrophobicite(position_med, hydrophobicite_moy):
    """ Fonction qui permet de créer un tableau des valeurs de l'hydrophobicité moyenne à partir de deux listes 
    input : position_med : liste avec les positions (entier) des acides aminés dans la séquence qui sont reliés à des profils d'hydrophobicité
            hydrophobicite_moy: liste avec les floattants des valeurs  d'hydrophobicité moyennes pour chaque fenetre glissante
    output : table des valeurs de l'hydrophobicité moyenne sous le format pandas"""
    df = pd.DataFrame({"Médiane": position_med, 'Fréquence Observée': hydrophobicite_moy })
    return df 

def graphique_profile_hydrophobicite(position_med, hydrophobicite_moyenne):
    """ Fonction qui cré une courbe du profil d'hydrophobicite en fonction des deux listes d'entrée 
    input : position_med : liste avec les positions (entier) des acides aminés dans la séquence qui sont reliés à des profils d'hydrophobicité
            hydrophobicite_moy: liste avec les floattants des valeurs  d'hydrophobicité moyennes pour chaque fenetre glissante
    output : graphique du proffil d'hydrophobicité matplotlib"""
    fig = Figure(figsize=(5, 4), dpi=100)
    sub = fig.add_subplot(111)
    sub.plot(position_med, hydrophobicite_moyenne)
    sub.set_xlabel('Position des acides aminés')
    sub.set_ylabel("Moyenne d'hydrophobicité")
    sub.set_title('Profil d\'hydrophobicité', fontweight = "bold")
    return fig

def raffraichissement_hydro(self,window_size):
    """Fonction qui permet d'afficher les nouvelles valeurs lorsque la valeur de la fenetre est changée 
    input : self = fenetre dans laquelle les valeurs sont affichées 
            window_size = valeur de la fenetre changée """
    position_med, hydrophobicite_moy= profil_hydrophobicite(sequence_1Le,window_size) #defini les nouvelles variables
    
    #RE DEFINI L'AFFICHAGE 
    #fenetre tableau
    tab1 = LabelFrame(self, text="Tableau de l'hydrophobicité moyenne", padx=2, pady=5)
    tab1.grid(row=6, column=0, sticky='nsew')
    tab_hydro = tableau_profile_hydrophobicite(position_med, hydrophobicite_moy)
    pt = Table(tab1, dataframe=tab_hydro, showtoolbar=True, showstatusbar=True)
    pt.show()

    #fenetre graphique
    graph= LabelFrame(self, text="Graphique de l'hydrophobicité moyenne", padx=2, pady=5)
    graph.grid(row=6, column=1, sticky='nsew')
    grap_hydro=graphique_profile_hydrophobicite(position_med, hydrophobicite_moy)
    canvas = FigureCanvasTkAgg(grap_hydro, master=graph)  # Création du canvas de Matplotlib
    canvas.draw()
    canvas.get_tk_widget().grid(row=7, column=1, sticky='nsew')


###FENETRE DE GESTION

#Utilisation du module tkinter pour permettre l'affichage des données générer 

class Application(Tk):
    def __init__(self):
        """Fonction qui permet l'initialisation de l'interface"""
        Tk.__init__(self)
        self._frame = None    
        self.switch_frame(StartPage)

        self.button_frame = tk.Frame(self)
        self.button_frame.pack(side="top")
        
        #definition des boutons qui serviront pour changer de fenetre 
        Button(self.button_frame, text="Description",
                  command=lambda: self.switch_frame(Page_description), height=2, width=12).grid(row=2, column=0, sticky='nsew')
        Button(self.button_frame, text=("Composition"+ "\n" +" en AA"),
                  command=lambda: self.switch_frame(Statistiques), height=2, width=12).grid(row=2, column=1, sticky='nsew')
        Button(self.button_frame, text=("Profil"+ "\n"+"Hydrophobicité"),
                      command=lambda: self.switch_frame(Profil_Hydro), height=2, width=12).grid(row=2, column=2, sticky='nsew')
        Button(self.button_frame, text="Physico"+ "\n"+"Chimie",
                  command=lambda: self.switch_frame(PhyChi), height=2, width=12).grid(row=2, column=4, sticky='nsew')
        Button(self.button_frame, text="Matrice"+ "\n"+"Contact",
                  command=lambda: self.switch_frame(Bonus), height=2, width=12).grid(row=2, column=5, sticky='nsew')
        Button(self.button_frame, text=("Nouvelle"+ "\n"+"Recherche"),
                      command=lambda: self.switch_frame(Nouvelle_recherche), height=2, width=12).grid(row=2, column=6, sticky='nsew')
        Button(self.button_frame, text=("Pont"+ "\n"+"disulfure"),
                      command=lambda: self.switch_frame(Structure), height=2, width=12).grid(row=2, column=3, sticky='nsew')

    def switch_frame(self, frame_class):
        """ Fonction qui permet de changer de fenetre 
        input : self= la fenetre où l'on affiche les données 
                frame_class= definition qui génère les valeurs à afficher"""
        new_frame = frame_class(self)
        if self._frame is not None: #si la fenetre existe déjà la détruire
            self._frame.destroy()
        self._frame = new_frame #afficher la nouvelle fenetre 
        self._frame.pack()

class StartPage(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Mot de bienvenu").pack(side="top", fill="x", pady=10) #titre de la fenetre 

        #definition du contenu à afficher sur la page 
        l = LabelFrame(self, text="Bienvenu à vous cher scientifique ! ", fg="blue",  padx=5, pady=5) #titre de la sous-fenetre 
        l.pack()
        script_start = (f" Vous allez accéder à une interface contenant les informations contenues dans la fiche PDB {Code_PDB.upper()}") #texte de la sous-fenetre 
        label_script = Label(l, text= script_start)
        label_script.pack() #afficher le contenu

class Page_description(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Description", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, columnspan=2, sticky='nsew') #titre 

        #sous-fenetre header
        d = LabelFrame(self, text="Header", fg="blue", padx=2, pady=5)
        d.grid(row=2, column=0, sticky='nsew')
        Script_d=affichage_informations()
        Label_Script = Label(d, text= Script_d, padx=2, pady=5) #localisation du texte
        Label_Script.grid(row=3, column=0, sticky='nsew') #affichage du texte 
        
        #sous-fenetre pour enregistrer les données 
        l = LabelFrame(self, text="Enregistrer les informations", fg="blue", padx=2, pady=5)
        l.grid(row=2, column=1, sticky='nsew')
        # au format fasta
        script_start = ("Cliquez pour ci-dessous pour enregistrer "+ "\n"+ "la fiche au format Fasta générée" + "\n"+ 
                        f"grâce aux informations de la fiche PDB {Code_PDB}. ") 
        label_script = Label(l, text= script_start)
        label_script.grid(row=3, column=1, sticky='nsew')
        button = Button(l, text="Enregistrer au format Fasta", command=lambda: fenetre_enr()) #Création d'un bouton pour exécuter les fonctions
        button.grid(row=4, column=1)
        # fichier réapitulatif 
        script_all= ("\n" +"Cliquez pour ci-dessous pour enregistrer "+ "\n"+ f"les données importantes de la fiche {Code_PDB}" + "\n"+ 
                        " ainsi que les analyses effectuées dans ces fenêtres." ) 
        label_script = Label(l, text= script_all)
        label_script.grid(row=6, column=1, sticky='nsew')
        button1 = Button(l, text="Enregistrer tout", command=lambda:enregistrer_final()) #Création d'un bouton pour exécuter les fonctions
        button1.grid(row=7, column=1)

class Statistiques(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Composition en acides aminés", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, columnspan=2, sticky='nsew')
        affichage_stat()#défini les variables nécéssaires
        
        #sous- fenetre informations 
        a = LabelFrame(self, text="Informations", fg="blue", padx=2, pady=5)
        a.grid(row=2, column=0, columnspan=2, sticky='nsew')
        script = (f"Sur cette page, vous pourrez visualiser les informations liées à la composition en acides aminées de la fiche au code PDB : {Code_PDB}." + "\n" +
                  "Vous pourrez aussi extraire le tableau des fréquences d'acides aminés en cliquant sur le bouton Enregistrer la tableau (en bas à gauche de la fenetre).")
        Label_Script = Label(a, text= script, padx=2, pady=5)
        Label_Script.grid(row=3, column=0,columnspan=2, sticky='nsew')

        #sosus- fenetre séquence 
        s = LabelFrame(self, text="Séquence d'acides aminés", fg="blue", padx=2, pady=5)
        s.grid(row=4, column=0, columnspan=2, sticky='nsew')
        Nbr= (f"La séquence contient {len(seq)} acides aminés." + "\n" + seq) 
        Nbr_script= Label(s, text= Nbr, padx=2, pady=5)
        Nbr_script.grid(row=5, column=0, sticky='nsew')
        
        #sous-fenetre tableau
        tab = LabelFrame(self, text="Tableau de fréquence", fg="blue", padx=2, pady=5)
        tab.grid(row=6, column=0, sticky='nsew')
        df=tableau_composition(Nombre_AA)
        # bouton pour enregistrer en xsls
        button = Button(tab, text="Enregistrer la tableau", command=lambda:enregistrer_tableau (df, "Composition_AA_"))
        button.grid(row=7, column=0 )
        pt = Table(tab, dataframe=df, showstatusbar=True)
        pt.show()

        #sous-fenetre graphique
        graph= LabelFrame(self, text="Graphique de fréquence",fg="blue", padx=2, pady=5)
        graph.grid(row=6, column=1, sticky='nsew')
        g=graphique_composition(Nombre_AA)
        canvas = FigureCanvasTkAgg(g, master=graph)  # Création du canvas de Matplotlib pour afficher le graphique
        canvas.draw()
        canvas.get_tk_widget().grid(row=7, column=1, sticky='nsew')

class PhyChi(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Physico-Chimie", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, sticky='nsew')
        
        #sous- fenetre caractéristiques
        p= LabelFrame(self, text="Caractéristiques", fg="blue", padx=2, pady=5)
        p.grid(row=2, column=0, sticky='nsew')
        script = ("Sur cette page vous pourrez télécharger un document pdb vous permettant de visualiser la structure étudiée avec différentes échelles de coloration." + 
                  "\n"*2 + "Pour cela, vous devez enregistrer le document en cliquant sur le bouton de votre choix : " + "\n " + 
                  "B-factor : visualisation selon l'agitation thermique " + "\n" + 
                  "Polarité : visualisation selon la polarité des acides aminés" + "\n" +
                  "Poids : visualisation selon le poids moléculaire des acides aminés" + "\n" + 
                  "Fréquence: visualisation selon la fréquence d'apparition dans la séquence" + "\n"*2  + 
                  "Ouvrez ensuite le fichier grâce à l'application PyMol, puis sélectionnez le mode de coloration 'Spectrum'. ")
        label_script = Label(p, text= script)
        label_script.grid(row=3, column=0, columnspan=4, sticky='nsew')
        
        #bouton pour télécharger le fichier
        button1 = Button(p, text="Polarité", height=2, width=10, command=lambda: fenetre_enr_hydro(2))
        button1.grid(row=4, column=1, sticky='nsew')
        button2 = Button(p, text="Poids", height=2, width=10, command=lambda: fenetre_enr_hydro(3))
        button2.grid(row=4, column=2, sticky='nsew')
        button3 = Button(p, text="Fréquence", height=2, width=10,command=lambda: fenetre_enr_hydro(4))
        button3.grid(row=4, column=3, sticky='nsew')
        button4 = Button(p, text="B-factor", height=2, width=10,command=lambda: fenetre_enr_hydro(1))
        button4.grid(row=4, column=0, sticky='nsew')
        
        #visualisation sur pymol du fichier lorsqu'il est téléchargé si possible
        try : 
            import pymol 
        except : 
            None 
        else : 
            #visualisation sur pymol
            pymol.finish_launching(['pymol', '-xi'])
            pymol.cmd.load(path_hydro)  # Chargez votre fichier PDB
            pymol.cmd.color('spectrum', path_hydro) 
            pymol.cmd.zoom()

class Bonus(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Matrice de Contact", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, sticky='nsew')
        #definition des variables nécessaires pour l'affichage des fonctions
        CALPHA = Calpha (data)
        matrice_contact = calcul_distance_calpha(CALPHA)
        
        #sous- fenetre matrice
        mat= LabelFrame(self, text="Matrice de contact", fg="blue", padx=2, pady=5)
        mat.grid(row=2, column=0, sticky='nsew')
        script =(f"Voici la matrice de contact de la protéine correspondante au code PDB {Code_PDB}." + "\n"*2 + 
                 "Une matrice de contact de protéine est une représentation matricielle des interactions entre les résidus"+ 
                 " d’une protéine tridimensionnelle." + "\n" +" Elle indique quels résidus sont en contact et quels résidus sont éloignés selon une échelle de coloration." + 
                 "\n" + " Elle permet de visualiser la structure secondaire et tertiaire de la protéine, ainsi que les domaines et les motifs structuraux.")
        script_lab = Label(mat, text= script)
        script_lab.grid(row=3, column=0, sticky='nsew')
        # afficher le graphique 
        m=afficher_matrice_contact(matrice_contact)
        canvas = FigureCanvasTkAgg(m, master=mat)  # Création du canvas de Matplotlib pour afficher le graphique
        canvas.draw()
        canvas.get_tk_widget().grid(row=4, column=0)

class Nouvelle_recherche(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Nouvelle Recherche", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, sticky='nsew')

        #sous- fenetre pour les instructions
        l = LabelFrame(self, text="Instructions", fg="blue",  padx=2, pady=5)
        l.grid(row=2, column=0, sticky='nsew')
        script_start = ("Pour relancer une nouvelle recherche à partir d'un autre code PDB ou d'un autre fichier local, cliquez sur Lancer." + 
                        "\n" + "Veuillez fermer cette fenêtre (en cliquant sur la croix rouge) après l'ouverture de la nouvelle pour éviter tout conflit de valeur.")
        label_script = Label(l, text= script_start)
        label_script.grid(row=3, column=0, sticky='nsew')
        button = Button(l, text="Lancer", command=lambda: initialisation())#Création d'un bouton pour exécuter les fonctions
        button.grid(row=4, column=0)

class Profil_Hydro(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Profil d'Hydrophobicité", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, columnspan=2, sticky='nsew')
        definition_var ()
        
        #sous-fenetre informations 
        a = LabelFrame(self, text="Informations",fg="blue",  padx=2, pady=5)
        a.grid(row=2, column=0, columnspan=2, sticky='nsew')
        script = (f"Sur cette page, vous pourrez visualiser les informations liées au profil d'hydrophobicité sous forme de tableau et de représentation graphique" + "\n"*2 +
                  "Le profil d'hydrophobicité des acides aminés de la séquence a été calculé grâce à l'échelle d'hydrophobicité de Fauchere et Pliska." + "\n" + 
                  "Il est calculé sur une fenêtre glissante de 9 acides aminés mais cette fenêtre est modifiable grâce au curseur ci-dessous." + "\n" + 
                  "Si vous souhaitez changer la taille de la fenêtre, veuillez faire glisser le curseur puis cliquer sur Changer le cadre. " + "\n"*2 + 
                  "Vous pourrez aussi extraire le tableau du profil d'hydrophobicité en cliquant sur le bouton Enregistrer la tableau (en bas à gauche de la fenetre).")
        Label_Script = Label(a, text= script, padx=2, pady=5)
        Label_Script.grid(row=3, column=0,columnspan=2,  sticky='nsew')
        
        #Création d'un scroll pour définir la valeur de la fenetre glissante
        window_size = DoubleVar(value=9) #definition d'une variable dépendante de tkinter avec comme valeur initiale 9
        button_scroll= Scale(a, variable=window_size, from_=1, to=(len(sequence_1Le)), orient=HORIZONTAL) #definition des bornes du scroll pour éviter des erreurs dans les fonctions suivantes
        button_scroll.grid(row=4, column=0,columnspan=2)
        button_val = Button(a, text="Changer le cadre", command=lambda: raffraichissement_hydro(self,window_size)) #Création d'un bouton pour exécuter les fonctions
        button_val.grid(row=5, column=0,columnspan=2)
        
        raffraichissement_hydro(self,window_size) #affichage des données grâce à la fonction

class Structure(tk.Frame):
    def __init__(self, master):
        tk.Frame.__init__(self, master)
        Label(self, text="Pont disulfure", fg="blue", font=("Helvetica",12, "bold")).grid(row=1, column=0, columnspan=2, sticky='nsew')
        
        #sous-fenetre Informations
        a = LabelFrame(self, text="Informations",fg="blue",  padx=2, pady=5)
        a.grid(row=2, column=0, columnspan=2, sticky='nsew')
        script = (f"Sur cette page, vous pourrez visualiser les informations liées à la structure décrite dans la fiche au code PDB {Code_PDB}."+ "\n" + 
                  "La présence éventuelle de ponts disulfures intramoléculaires a été étudié pour cette fiche et les résultats sont répertoriés ci-dessous." + 
                  "\n"*2 + "NB : si les cystines présentes ne sont pas impliqués dans la structure elle pourrait participer à un pont disulfure avec une autre molécule (non étudiée ici).")
        Label_Script = Label(a, text= script, padx=2, pady=5)
        Label_Script.grid(row=3, column=0, sticky='nsew')
        
        #Sous- fenetre Pont disulfure 
        b = LabelFrame(self, text="Pont disulfure",fg="blue",  padx=2, pady=5)
        b.grid(row=5, column=0, columnspan=2, sticky='nsew') 
        Label_Script = Label(b, text= affichage_stru(), padx=2, pady=5)
        Label_Script.grid(row=6, column=0, sticky='nsew')


##====MAIN=====
initialisation() 
